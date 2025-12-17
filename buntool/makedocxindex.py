from dataclasses import dataclass

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.styles.style import CharacterStyle

from buntool.headers import HEADERS


@dataclass
class DocxConfig:
    """Configuration for DOCX Table of Contents generation."""

    confidential: bool = False
    date_setting: bool = True
    index_font_setting: str | None = None


def _setup_document_style(doc: DocumentObject, index_font_setting):
    style = doc.styles["Normal"]
    if isinstance(style, CharacterStyle):
        font_name = "Times New Roman"
        if index_font_setting == "sans":
            font_name = "Arial"
        elif index_font_setting == "serif":
            font_name = "Times New Roman"
        elif index_font_setting == "mono":
            font_name = "Courier New"

        style.font.name = font_name


def _add_docx_header(doc: DocumentObject, casedetails: dict[str, str], confidential: bool):
    # Set up case details
    claimno_hdr = casedetails.get("claim_no", "")
    casename = casedetails.get("case_name", "")
    bundle_name = casedetails.get("bundle_title", "").upper()

    # Add the Claim Number (if exists)
    if claimno_hdr:
        para = doc.add_paragraph(claimno_hdr)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add Case Name (if exists)
    if casename:
        para = doc.add_paragraph(casename)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(14)

    # Add the Bundle Name
    if bundle_name:
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(bundle_name)
        run.bold = True
        run.font.size = Pt(16)
        if confidential:
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.text = f"CONFIDENTIAL\n{bundle_name}"


def _create_and_populate_table(doc: DocumentObject, toc_entries, date_setting):
    # Table of Contents header row
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = HEADERS[0]
    header_cells[1].text = HEADERS[1]
    header_cells[2].text = HEADERS[2] if date_setting else ""
    header_cells[3].text = HEADERS[3]
    for cell in header_cells:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    # Add entries to the Table of Contents
    for entry in toc_entries:
        # Skip the header row if it's present in the data
        if "Tab" in entry[0] and "Title" in entry[1]:
            continue

        row = table.add_row().cells
        if "SECTION_BREAK" in entry[0]:
            # Handle section breaks
            row[0].merge(row[-1])  # Merge all cells for a section header
            para = row[0].paragraphs[0]
            run = para.add_run(entry[1])
            run.bold = True
            run.font.size = Pt(12)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            # Add a regular TOC entry
            row[0].text = str(entry[0])  # Tab
            row[1].text = entry[1]  # Title
            row[2].text = entry[2] if date_setting else ""  # Date
            row[3].text = str(entry[3])  # Page


def create_toc_docx(toc_entries, casedetails: dict[str, str], output_file_path, config: DocxConfig):
    """Creates a Table of Contents in a .docx file."""
    doc = Document()

    _setup_document_style(doc, config.index_font_setting)
    _add_docx_header(doc, casedetails, config.confidential)
    _create_and_populate_table(doc, toc_entries, config.date_setting)

    # Save the document
    doc.save(output_file_path)


# main
if __name__ == "__main__":
    # Sample data
    toc_entries = [
        ("001.", "First Doc", "2021-01-01", 1),
        ("002", "Second Doc", "2021-01-02", 5),
        ("003", "Third Document", "2021-01-03", 10),
        # ("SECTION_BREAK", "Section Break Test", "", ""),
        ("004", "Document Number Four", "2021-01-04", 15),
        ("005", "The fifth document in this series", "2021-01-05", 20),
    ]
    casedetails = {"bundle_title": "Bundle Name", "claim_no": "Claim Number", "case_name": "Case Name"}
    output_file_path = "TOC.docx"

    docx_config = DocxConfig(confidential=True, date_setting=False)
    # Create the TOC document
    create_toc_docx(toc_entries, casedetails, output_file_path, docx_config)
